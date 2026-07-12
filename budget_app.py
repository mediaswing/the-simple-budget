#!/usr/bin/env python3
"""Budget tracker GUI.

A self-contained Tkinter application that stores data in its own SQLite
database (``budget.db``). The schema mirrors ``budget.sql`` (a MariaDB dump):

  * ``budget_items`` -- id, category (the spending categories)
  * ``budget_lines`` -- id, store, price, date, category (FK -> budget_items)

The window has two tabs:
  * "Data"  -- add / edit / delete budget lines.
  * "Chart" -- a column graph of total spend (GBP) grouped by category,
               year, or three-letter month name.

The bar chart is drawn directly on a Tkinter Canvas, so the only requirement
is a standard Python install with Tk (no third-party packages needed).
"""

import os
import io
import sys
import csv
import queue
import sqlite3
import threading
import subprocess
import configparser
import datetime as dt
import tkinter as tk
import webbrowser
from tkinter import ttk, messagebox, filedialog

import updater

__version__ = "0.6.2"

# pyttsx3 powers the "read chart aloud" feature for visually impaired users.
# It is optional: if it is not installed the app still runs, and the speech
# controls explain how to enable it.
try:
    import pyttsx3
except ImportError:  # pragma: no cover - depends on the environment
    pyttsx3 = None

# pymysql is only needed when the user points budget.ini at a central MariaDB
# server. It is optional: without it (or without a config file) the app uses its
# own local SQLite database.
try:
    import pymysql
    import pymysql.cursors
except ImportError:  # pragma: no cover - depends on the environment
    pymysql = None

# Optional report-export backends. Each format is available only if its library
# is installed; the Report tab disables the ones that are missing and explains
# how to enable them.
try:
    from fpdf import FPDF
except ImportError:  # pragma: no cover - depends on the environment
    FPDF = None

try:
    import docx
except ImportError:  # pragma: no cover - depends on the environment
    docx = None

# Duplicate-key errors differ per backend; catch both to give the same message.
INTEGRITY_ERRORS = (sqlite3.IntegrityError,) + (
    (pymysql.err.IntegrityError,) if pymysql is not None else ())

HERE = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(HERE, "budget.db")
INI_PATH = os.path.join(HERE, "budget.ini")

MONTHS = ["Jan", "Feb", "Mar", "Apr", "May", "Jun",
          "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]

# Categories seeded from budget.sql (id -> category). Ids are preserved so the
# SQLite database matches the original MariaDB data.
SEED_CATEGORIES = [
    (20, "Household Stuff"),
    (21, "Food/Drink"),
    (22, "Utilities/Bills"),
    (23, "Entertainment"),
    (25, "Web Services"),
    (28, "Healthcare"),
    (31, "Transport"),
    (32, "Accommodation"),
]


def gbp(value):
    """Format a number as pounds sterling, e.g. 1234.5 -> '£1,234.50'."""
    sign = "-" if value < 0 else ""
    return "{}£{:,.2f}".format(sign, abs(value))


def spoken_gbp(value):
    """Format a number as words a screen reader / TTS can pronounce clearly.

    e.g. 1234.5 -> '1234 pounds and 50 pence'.
    """
    sign = "minus " if value < 0 else ""
    value = abs(value)
    pounds = int(value)
    pence = int(round((value - pounds) * 100))
    if pence == 100:  # rounding pushed us to the next pound
        pounds += 1
        pence = 0
    text = "{}{} pound{}".format(sign, pounds, "" if pounds == 1 else "s")
    if pence:
        text += " and {} pence".format(pence)
    return text


# --------------------------------------------------------------------------- #
# Text to speech
# --------------------------------------------------------------------------- #
class Speaker:
    """Thin, thread-safe wrapper around pyttsx3.

    Speech runs on a background thread so the GUI stays responsive. Starting a
    new utterance stops any that is in progress. All methods are safe to call
    even when pyttsx3 is not installed (they simply do nothing).
    """

    def __init__(self):
        self._lock = threading.Lock()
        self._engine = None
        self._thread = None
        # User-adjustable settings, applied to each utterance. None means
        # "leave the engine default alone".
        self.rate = None
        self.voice_id = None
        self._voices_cache = None
        self._voices_lock = threading.Lock()
        self._default_rate = 200
        self._default_voice = None
        # pyttsx3.init() (in voices(), below) can take a noticeable fraction
        # of a second -- start probing it now, on a background thread, so it
        # has a head start on the rest of BudgetApp's construction instead of
        # blocking first paint when the Chart tab asks for the voice list a
        # moment later. Safe on a non-main thread: speak() (below) already
        # runs its own pyttsx3.init() off the main thread.
        threading.Thread(target=self.voices, daemon=True).start()

    @property
    def available(self):
        return pyttsx3 is not None

    @property
    def default_rate(self):
        self.voices()  # ensure the engine defaults have been probed
        return self._default_rate

    def voices(self):
        """Return [(voice_id, name), ...] for the installed voices (cached).

        Thread-safe: the initial probe may be kicked off from __init__'s
        background thread and/or called again from the main thread before
        that finishes, so only one of them should actually call pyttsx3.init().
        """
        if not self.available:
            return []
        if self._voices_cache is None:
            with self._voices_lock:
                if self._voices_cache is None:  # re-check: lost the race?
                    self._voices_cache = []
                    try:
                        engine = pyttsx3.init()
                        self._voices_cache = [
                            (v.id, v.name) for v in engine.getProperty("voices")]
                        self._default_rate = engine.getProperty("rate") or 200
                        self._default_voice = engine.getProperty("voice")
                        engine.stop()
                    except Exception:  # pragma: no cover - driver issues
                        pass
        return self._voices_cache

    def speak(self, text):
        """Speak *text* aloud. Returns True if speech was started."""
        if not self.available or not text:
            return False
        self.stop()
        self._thread = threading.Thread(
            target=self._run, args=(text,), daemon=True)
        self._thread.start()
        return True

    def _run(self, text):
        with self._lock:
            try:
                engine = pyttsx3.init()
                self._engine = engine
                if self.rate is not None:
                    engine.setProperty("rate", self.rate)
                if self.voice_id is not None:
                    engine.setProperty("voice", self.voice_id)
                engine.say(text)
                engine.runAndWait()
            except Exception:  # pragma: no cover - engine/driver issues
                pass
            finally:
                self._engine = None

    def stop(self):
        """Interrupt any speech that is currently playing."""
        engine = self._engine
        if engine is not None:
            try:
                engine.stop()
            except Exception:  # pragma: no cover
                pass

    def save_to_file(self, text, path, on_done):
        """Render *text* to an audio file on a background thread.

        *on_done* is called with None on success or the exception on failure.
        It runs on the worker thread, so it should only hand results back to the
        GUI thread in a thread-safe way (e.g. via a queue).
        """
        if not self.available:
            return False

        def run():
            err = None
            try:
                engine = pyttsx3.init()
                if self.rate is not None:
                    engine.setProperty("rate", self.rate)
                if self.voice_id is not None:
                    engine.setProperty("voice", self.voice_id)
                engine.save_to_file(text, path)
                engine.runAndWait()
            except Exception as e:  # pragma: no cover - driver issues
                err = e
            on_done(err)

        threading.Thread(target=run, daemon=True).start()
        return True


# --------------------------------------------------------------------------- #
# Database configuration
# --------------------------------------------------------------------------- #
class DBConfig:
    """Resolved database choice: which backend and how to connect to it."""

    def __init__(self, backend, params, description):
        self.backend = backend          # "sqlite" or "mariadb"
        self.params = params            # dict of connection parameters
        self.description = description  # human-readable, for the UI


def load_db_config(ini_path=INI_PATH):
    """Decide which database to use, based on an optional INI file.

    The file (default ``budget.ini`` next to the app) may contain::

        [database]
        backend = mariadb          ; optional: "sqlite" or "mariadb"
        host = db.example.com
        port = 3306
        user = budget
        password = secret
        database = budget

    If ``backend`` is omitted the choice is inferred: MariaDB is used when
    host, user, and database are all provided, otherwise the internal SQLite
    database. With no config file at all, SQLite is always used.
    """
    parser = configparser.ConfigParser()
    if not parser.read(ini_path) or not parser.has_section("database"):
        return DBConfig("sqlite", {"path": DB_PATH}, "internal database (SQLite)")

    sec = parser["database"]
    backend = sec.get("backend", "").strip().lower()
    host = sec.get("host", "").strip()
    user = sec.get("user", "").strip()
    database = sec.get("database", "").strip()
    password = sec.get("password", "")
    port = sec.getint("port", 3306)
    has_credentials = bool(host and user and database)

    want_mariadb = backend in ("mariadb", "mysql") or (
        not backend and has_credentials)

    if want_mariadb:
        if not has_credentials:
            raise ValueError(
                "budget.ini requests MariaDB but is missing host/user/database.")
        params = dict(host=host, port=port, user=user,
                      password=password, database=database)
        return DBConfig("mariadb", params,
                        "MariaDB ({}/{})".format(host, database))

    path = sec.get("path", "").strip() or DB_PATH
    return DBConfig("sqlite", {"path": path}, "internal database (SQLite)")


def open_db(ini_path=INI_PATH):
    """Open the configured database, falling back to SQLite on any problem.

    Returns ``(db, warning)`` where *warning* is None, or a message describing
    why the app fell back to the internal database (shown to the user).
    """
    sqlite_cfg = DBConfig("sqlite", {"path": DB_PATH}, "internal database (SQLite)")

    try:
        cfg = load_db_config(ini_path)
    except Exception as e:  # malformed budget.ini (bad port, missing fields, ...)
        return (BudgetDB(sqlite_cfg),
                "Could not read budget.ini: {}\n\n"
                "Using the internal database instead.".format(e))

    if cfg.backend != "mariadb":
        try:
            return BudgetDB(cfg), None
        except Exception as e:  # bad/unwritable configured path
            return (BudgetDB(sqlite_cfg),
                    "Could not open the configured database file ({}): {}\n\n"
                    "Using the internal database instead.".format(
                        cfg.params.get("path"), e))

    if pymysql is None:
        return (BudgetDB(sqlite_cfg),
                "budget.ini requests MariaDB but the 'pymysql' package is not "
                "installed. Using the internal database instead.\n\n"
                "Install it with:  pip install pymysql")
    try:
        return BudgetDB(cfg), None
    except Exception as e:  # connection / auth / schema errors
        return (BudgetDB(sqlite_cfg),
                "Could not connect to MariaDB ({}): {}\n\n"
                "Using the internal database instead.".format(
                    cfg.description, e))


# --------------------------------------------------------------------------- #
# Access control (optional; Windows AD domain / Microsoft Entra)
# --------------------------------------------------------------------------- #
# Keeps the helper subprocesses from flashing a console window when the app runs
# as a windowed (console=False) Windows binary. Empty (no-op) off Windows, where
# CREATE_NO_WINDOW doesn't exist.
_NO_WINDOW = ({"creationflags": subprocess.CREATE_NO_WINDOW}
              if sys.platform == "win32" else {})


class AccessPolicy:
    """Who may run the app: membership of a named Windows security group."""

    def __init__(self, group, deny_on_error=True):
        self.group = group                  # required group (name or SID)
        self.deny_on_error = deny_on_error  # fail closed if membership unknown


def load_access_config(ini_path=INI_PATH):
    """Read the optional ``[access]`` section of budget.ini.

    Returns an :class:`AccessPolicy` when a ``group`` is configured, else
    ``None`` (unrestricted)::

        [access]
        group = Budget-Users     ; required security group (name or SID)
        deny_on_error = true     ; optional; deny if membership can't be checked
    """
    parser = configparser.ConfigParser()
    if not parser.read(ini_path) or not parser.has_section("access"):
        return None
    sec = parser["access"]
    group = sec.get("group", "").strip()
    if not group:
        return None
    return AccessPolicy(group, sec.getboolean("deny_on_error", True))


def _run_windows_tool(args, timeout=10):
    """Run a Windows diagnostic command, returning its CompletedProcess if it
    ran, exited zero, and produced output -- or None otherwise (missing
    binary, timeout, non-zero exit, or empty stdout)."""
    try:
        out = subprocess.run(args, capture_output=True, text=True,
                             timeout=timeout, **_NO_WINDOW)
    except (OSError, subprocess.SubprocessError):
        return None
    if out.returncode != 0 or not out.stdout:
        return None
    return out


def windows_join_status():
    """Return ``(domain_joined, entra_joined)`` for this Windows machine, or
    ``None`` if the join state could not be determined.

    Uses ``dsregcmd /status`` (Windows 10+), falling back to an AD-domain logon
    heuristic. ``(False, False)`` means "determined: standalone"; ``None`` means
    the check itself failed (dsregcmd missing/blocked/timed out) so the caller
    can decide whether to fail open or closed.
    """
    out = _run_windows_tool(["dsregcmd", "/status"])
    if out is not None:
        domain = entra = False
        for line in out.stdout.splitlines():
            low = line.lower().replace(" ", "")
            if "azureadjoined" in low:
                entra = "yes" in low
            elif "domainjoined" in low:
                domain = "yes" in low
        return domain, entra
    # dsregcmd unavailable: a domain logon still exposes the DNS domain name.
    if os.environ.get("USERDNSDOMAIN"):
        return True, False
    return None  # genuinely undeterminable


def current_user_groups():
    """Set of the current user's group identities (names and SIDs, lowercased),
    or ``None`` if membership could not be determined.

    Uses ``whoami /groups``; group names appear as ``DOMAIN\\Group`` and are
    also stored bare (``group``). SIDs are included so an Entra cloud group can
    be matched by its object SID when its name doesn't resolve locally.
    """
    out = _run_windows_tool(["whoami", "/groups", "/fo", "csv", "/nh"])
    if out is None:
        return None
    ids = set()
    for row in csv.reader(io.StringIO(out.stdout)):
        if not row:
            continue
        name = row[0].strip()
        if name:
            ids.add(name.lower())
            if "\\" in name:
                ids.add(name.split("\\", 1)[1].lower())
        if len(row) > 2 and row[2].strip():
            ids.add(row[2].strip().lower())
    return ids or None


def access_denied_reason(policy):
    """Return a message explaining why access is denied, or ``None`` to allow.

    The group check applies only on a Windows machine that is joined to an AD
    domain or to Microsoft Entra; anywhere else there is no directory to check
    against, so access is allowed.
    """
    if policy is None or not policy.group:
        return None
    if sys.platform != "win32":
        return None
    unverified = ("Could not verify your membership of the security group "
                  "'{}' required to use this application. Access is denied."
                  .format(policy.group))
    status = windows_join_status()
    if status is None:
        # Couldn't tell whether the machine is joined; treat as unverifiable.
        return unverified if policy.deny_on_error else None
    if not any(status):
        return None  # determined standalone: no directory to enforce against
    groups = current_user_groups()
    if groups is None:
        return unverified if policy.deny_on_error else None
    want = policy.group.lower()
    if want in groups or ("\\" in want and want.split("\\", 1)[1] in groups):
        return None
    return ("Access denied: your account is not a member of the security group "
            "'{}' required to use this application.\n\nContact your "
            "administrator if you believe this is an error.".format(policy.group))


# --------------------------------------------------------------------------- #
# Data access
# --------------------------------------------------------------------------- #
class BudgetDB:
    """Backend-agnostic wrapper over SQLite or MariaDB (same public API)."""

    def __init__(self, config=None):
        self.config = config or DBConfig(
            "sqlite", {"path": DB_PATH}, "internal database (SQLite)")
        self.backend = self.config.backend
        if self.backend == "mariadb":
            self._ph = "%s"
            self.conn = self._connect_mariadb()
        else:
            self._ph = "?"
            self.conn = sqlite3.connect(self.config.params["path"])
            self.conn.row_factory = sqlite3.Row
            self.conn.execute("PRAGMA foreign_keys = ON")
        self._create_schema()
        self._seed_categories()

    # -- connection --------------------------------------------------------- #
    def _connect_mariadb(self):
        """Connect to MariaDB, creating the target database if it is missing.

        MariaDB rejects the connection with error 1049 ("Unknown database")
        when the configured schema does not exist yet. In that case we connect
        to the server without selecting a database, issue ``CREATE DATABASE IF
        NOT EXISTS``, and retry -- so a fresh server only needs a user with the
        CREATE privilege, not a pre-created schema.
        """
        try:
            return self._open_mariadb_connection()
        except pymysql.err.OperationalError as e:
            if e.args and e.args[0] == 1049:  # Unknown database
                self._create_database()
                return self._open_mariadb_connection()
            raise

    def _open_mariadb_connection(self):
        return pymysql.connect(
            cursorclass=pymysql.cursors.DictCursor,
            charset="utf8mb4", autocommit=False, **self.config.params)

    def _create_database(self):
        """Create the configured MariaDB schema (used on first connect)."""
        name = self.config.params["database"]
        server = {k: v for k, v in self.config.params.items() if k != "database"}
        conn = pymysql.connect(charset="utf8mb4", autocommit=True, **server)
        try:
            with conn.cursor() as cur:
                # Quote the identifier, escaping any embedded backticks.
                cur.execute("CREATE DATABASE IF NOT EXISTS `{}` "
                            "CHARACTER SET utf8mb4".format(name.replace("`", "``")))
        finally:
            conn.close()

    # -- low-level helpers -------------------------------------------------- #
    def _query(self, sql, params=None):
        """Run a SELECT and return a list of dict-like rows."""
        if self.backend == "mariadb":
            with self.conn.cursor() as cur:
                cur.execute(sql, params)  # None => no % substitution
                return cur.fetchall()
        return self.conn.execute(sql, params or ()).fetchall()

    def _exec(self, sql, params=None):
        """Run a write statement and commit."""
        if self.backend == "mariadb":
            with self.conn.cursor() as cur:
                cur.execute(sql, params)
        else:
            self.conn.execute(sql, params or ())
        self.conn.commit()

    def _create_schema(self):
        if self.backend == "mariadb":
            items = (
                "CREATE TABLE IF NOT EXISTS budget_items ("
                " id INT PRIMARY KEY AUTO_INCREMENT,"
                " category VARCHAR(255) NOT NULL UNIQUE)")
            lines = (
                "CREATE TABLE IF NOT EXISTS budget_lines ("
                " id INT PRIMARY KEY AUTO_INCREMENT,"
                " store VARCHAR(255) NOT NULL,"
                " price DECIMAL(10,2) NOT NULL,"
                " date DATE NOT NULL,"
                " category INT NOT NULL,"
                " FOREIGN KEY (category) REFERENCES budget_items(id)"
                " ON DELETE CASCADE)")
        else:
            items = (
                "CREATE TABLE IF NOT EXISTS budget_items ("
                " id INTEGER PRIMARY KEY,"
                " category TEXT NOT NULL UNIQUE)")
            lines = (
                "CREATE TABLE IF NOT EXISTS budget_lines ("
                " id INTEGER PRIMARY KEY AUTOINCREMENT,"
                " store TEXT NOT NULL,"
                " price REAL NOT NULL,"
                " date TEXT NOT NULL,"
                " category INTEGER NOT NULL,"
                " FOREIGN KEY (category) REFERENCES budget_items(id)"
                " ON DELETE CASCADE)")
        self._exec(items)
        self._exec(lines)
        if self.backend == "mariadb":
            self._ensure_autoincrement()

    def _ensure_autoincrement(self):
        """Backfill AUTO_INCREMENT on budget_items.id for MariaDB databases
        created by an earlier version (before categories could be added).

        ``CREATE TABLE IF NOT EXISTS`` leaves an existing table untouched, so a
        v0.2.0 database keeps its plain ``INT PRIMARY KEY`` and can't accept a
        category insert without an explicit id. Detect that and fix it in place.
        """
        rows = self._query(
            "SELECT EXTRA FROM information_schema.COLUMNS "
            "WHERE TABLE_SCHEMA = DATABASE() AND TABLE_NAME = 'budget_items' "
            "AND COLUMN_NAME = 'id'")
        extra = (rows[0]["EXTRA"] if rows else "") or ""
        if "auto_increment" not in extra.lower():
            self._exec(
                "ALTER TABLE budget_items MODIFY id INT NOT NULL AUTO_INCREMENT")

    def _seed_categories(self):
        verb = "INSERT IGNORE" if self.backend == "mariadb" else "INSERT OR IGNORE"
        sql = "{} INTO budget_items (id, category) VALUES ({}, {})".format(
            verb, self._ph, self._ph)
        for cid, name in SEED_CATEGORIES:
            self._exec(sql, (cid, name))

    # -- categories --------------------------------------------------------- #
    def categories(self):
        """Return list of (id, category) ordered by name."""
        rows = self._query(
            "SELECT id, category FROM budget_items ORDER BY category")
        return [(r["id"], r["category"]) for r in rows]

    def categories_with_stats(self):
        """Return [{id, name, lines, total}, ...] for the category manager."""
        rows = self._query(
            """
            SELECT i.id AS id, i.category AS name,
                   COUNT(l.id) AS lines, COALESCE(SUM(l.price), 0) AS total
            FROM budget_items i
            LEFT JOIN budget_lines l ON l.category = i.id
            GROUP BY i.id, i.category
            ORDER BY i.category
            """
        )
        return [
            {"id": r["id"], "name": r["name"],
             "lines": int(r["lines"]), "total": float(r["total"] or 0)}
            for r in rows
        ]

    def add_category(self, name):
        """Create a new category. Raises ValueError if the name is taken."""
        try:
            self._exec(
                "INSERT INTO budget_items (category) VALUES ({})".format(self._ph),
                (name,))
        except INTEGRITY_ERRORS:
            raise ValueError("A category named '{}' already exists.".format(name))

    def rename_category(self, cat_id, name):
        """Rename a category. Raises ValueError if the new name is taken."""
        try:
            self._exec(
                "UPDATE budget_items SET category={0} WHERE id={0}".format(self._ph),
                (name, cat_id))
        except INTEGRITY_ERRORS:
            raise ValueError("A category named '{}' already exists.".format(name))

    def delete_category(self, cat_id):
        """Delete a category. Refuses if any budget lines still use it."""
        n = self._query(
            "SELECT COUNT(*) AS n FROM budget_lines WHERE category={}".format(
                self._ph),
            (cat_id,))[0]["n"]
        if n:
            raise ValueError(
                "This category is used by {} budget line{}. Reassign or delete "
                "those first.".format(n, "" if n == 1 else "s"))
        self._exec("DELETE FROM budget_items WHERE id={}".format(self._ph),
                   (cat_id,))

    # -- lines -------------------------------------------------------------- #
    def lines(self):
        """Return all budget lines joined with their category name.

        Rows are normalised to plain dicts so the two backends look identical
        to callers (notably: dates as ISO strings, prices as floats).
        """
        rows = self._query(
            """
            SELECT l.id, l.store, l.price, l.date,
                   l.category AS category_id, i.category AS category_name
            FROM budget_lines l
            JOIN budget_items i ON i.id = l.category
            ORDER BY l.date DESC, l.id DESC
            """
        )
        return [
            {"id": r["id"], "store": r["store"],
             "price": float(r["price"]), "date": str(r["date"])[:10],
             "category_id": r["category_id"],
             "category_name": r["category_name"]}
            for r in rows
        ]

    def add_line(self, store, price, date, category_id):
        cols = ", ".join([self._ph] * 4)
        self._exec(
            "INSERT INTO budget_lines (store, price, date, category) "
            "VALUES ({})".format(cols),
            (store, price, date, category_id))

    def update_line(self, line_id, store, price, date, category_id):
        ph = self._ph
        self._exec(
            "UPDATE budget_lines SET store={0}, price={0}, date={0}, "
            "category={0} WHERE id={0}".format(ph),
            (store, price, date, category_id, line_id))

    def delete_line(self, line_id):
        self._exec("DELETE FROM budget_lines WHERE id={}".format(self._ph),
                   (line_id,))

    # -- aggregation for the chart ----------------------------------------- #
    def totals_by(self, group):
        """Return [(label, total_price), ...] grouped as requested.

        group is one of 'category', 'year', 'month'.
        """
        maria = self.backend == "mariadb"

        if group == "category":
            rows = self._query(
                """
                SELECT i.category AS label, COALESCE(SUM(l.price), 0) AS total
                FROM budget_lines l
                JOIN budget_items i ON i.id = l.category
                GROUP BY i.id, i.category
                ORDER BY total DESC
                """
            )
            return [(r["label"], float(r["total"] or 0)) for r in rows]

        if group == "year":
            year = "DATE_FORMAT(date, '%Y')" if maria else "strftime('%Y', date)"
            rows = self._query(
                "SELECT {} AS label, SUM(price) AS total "
                "FROM budget_lines GROUP BY label ORDER BY label".format(year)
            )
            return [(r["label"], float(r["total"] or 0)) for r in rows]

        if group == "month":
            mon = "DATE_FORMAT(date, '%m')" if maria else "strftime('%m', date)"
            rows = self._query(
                "SELECT {} AS mnum, SUM(price) AS total "
                "FROM budget_lines GROUP BY mnum".format(mon)
            )
            by_month = {r["mnum"]: float(r["total"] or 0) for r in rows}
            # Show all 12 months in calendar order (0 where there's no spend).
            return [
                (MONTHS[i], by_month.get("{:02d}".format(i + 1), 0))
                for i in range(12)
            ]

        raise ValueError("unknown group: %r" % group)


# --------------------------------------------------------------------------- #
# Categories tab
# --------------------------------------------------------------------------- #
class CategoriesTab(ttk.Frame):
    """Add / rename / delete the user's own budget categories."""

    def __init__(self, master, db, on_change):
        super().__init__(master, padding=10)
        self.db = db
        self.on_change = on_change  # called after any change (refresh siblings)
        self.selected_id = None

        self._build_form()
        self._build_table()
        self._bind_keys()
        self.refresh()

    def _build_form(self):
        form = ttk.LabelFrame(self, text="Category", padding=10)
        form.pack(fill="x")

        ttk.Label(form, text="Name").grid(row=0, column=0, sticky="w",
                                          padx=4, pady=4)
        self.name_var = tk.StringVar()
        self.name_entry = ttk.Entry(form, textvariable=self.name_var, width=30)
        self.name_entry.grid(row=0, column=1, padx=4, pady=4)

        btns = ttk.Frame(form)
        btns.grid(row=1, column=0, columnspan=2, sticky="w", pady=(8, 0))
        self.add_btn = ttk.Button(btns, text="Add (Ctrl+S)", command=self.add,
                                  underline=0)
        self.add_btn.pack(side="left", padx=(0, 6))
        self.rename_btn = ttk.Button(
            btns, text="Rename selected (Ctrl+S)", command=self.rename,
            state="disabled")
        self.rename_btn.pack(side="left", padx=6)
        self.delete_btn = ttk.Button(
            btns, text="Delete selected (Del)", command=self.delete,
            state="disabled")
        self.delete_btn.pack(side="left", padx=6)
        ttk.Button(btns, text="Clear (Esc)", command=self.clear_form).pack(
            side="left", padx=6)

    def _build_table(self):
        wrap = ttk.Frame(self)
        wrap.pack(fill="both", expand=True, pady=(10, 0))

        cols = ("name", "lines", "total")
        self.tree = ttk.Treeview(wrap, columns=cols, show="headings",
                                 selectmode="browse")
        headings = {"name": "Category", "lines": "Lines", "total": "Total spend"}
        widths = {"name": 240, "lines": 80, "total": 140}
        for c in cols:
            self.tree.heading(c, text=headings[c])
            anchor = "w" if c == "name" else "e"
            self.tree.column(c, width=widths[c], anchor=anchor)
        self.tree.pack(side="left", fill="both", expand=True)

        sb = ttk.Scrollbar(wrap, orient="vertical", command=self.tree.yview)
        sb.pack(side="right", fill="y")
        self.tree.configure(yscrollcommand=sb.set)
        self.tree.bind("<<TreeviewSelect>>", self.on_select)
        self.tree.bind("<Delete>", lambda _e: self.delete())

    def _bind_keys(self):
        def _clear(_e):
            self.clear_form()
            return "break"

        self.name_entry.bind("<Return>", lambda _e: self.save())
        self.name_entry.bind("<Escape>", _clear)

    # -- helpers ------------------------------------------------------------ #
    def refresh(self):
        for item in self.tree.get_children():
            self.tree.delete(item)
        for row in self.db.categories_with_stats():
            self.tree.insert(
                "", "end", iid=str(row["id"]),
                values=(row["name"], row["lines"], gbp(row["total"])))

    # -- actions ------------------------------------------------------------ #
    def save(self):
        """Rename the selected category, or add a new one (Ctrl+S / Enter)."""
        if self.selected_id is None:
            self.add()
        else:
            self.rename()

    def focus_new(self):
        self.clear_form()
        self.name_entry.focus_set()

    def _read_name(self):
        name = self.name_var.get().strip()
        if not name:
            raise ValueError("Please enter a category name.")
        return name

    def add(self):
        try:
            name = self._read_name()
            self.db.add_category(name)
        except ValueError as e:
            messagebox.showerror("Invalid input", str(e))
            return
        self.clear_form()
        self.refresh()
        self.on_change()

    def rename(self):
        if self.selected_id is None:
            return
        try:
            name = self._read_name()
            self.db.rename_category(self.selected_id, name)
        except ValueError as e:
            messagebox.showerror("Invalid input", str(e))
            return
        self.clear_form()
        self.refresh()
        self.on_change()

    def delete(self):
        if self.selected_id is None:
            return
        name = self.name_var.get().strip()
        if not messagebox.askyesno(
                "Delete", "Delete the category '{}'?".format(name)):
            return
        try:
            self.db.delete_category(self.selected_id)
        except ValueError as e:
            messagebox.showerror("Cannot delete", str(e))
            return
        self.clear_form()
        self.refresh()
        self.on_change()

    def on_select(self, _event=None):
        sel = self.tree.selection()
        if not sel:
            return
        self.selected_id = int(sel[0])
        self.name_var.set(self.tree.item(sel[0], "values")[0])
        self.rename_btn.config(state="normal")
        self.delete_btn.config(state="normal")

    def clear_form(self):
        self.selected_id = None
        if self.tree.selection():
            self.tree.selection_remove(self.tree.selection())
        self.name_var.set("")
        self.rename_btn.config(state="disabled")
        self.delete_btn.config(state="disabled")


# --------------------------------------------------------------------------- #
# Data tab
# --------------------------------------------------------------------------- #
class DataTab(ttk.Frame):
    """Add / edit / delete budget lines."""

    def __init__(self, master, db, on_change):
        super().__init__(master, padding=10)
        self.db = db
        self.on_change = on_change
        self.selected_id = None
        self._cat_by_name = {}

        self._build_form()
        self._build_table()
        self._bind_keys()
        self.refresh()

    def _build_form(self):
        form = ttk.LabelFrame(self, text="Budget line", padding=10)
        form.pack(fill="x")

        ttk.Label(form, text="Store").grid(row=0, column=0, sticky="w", padx=4, pady=4)
        self.store_var = tk.StringVar()
        self.store_entry = ttk.Entry(form, textvariable=self.store_var, width=24)
        self.store_entry.grid(row=0, column=1, padx=4, pady=4)

        ttk.Label(form, text="Price (£)").grid(row=0, column=2, sticky="w", padx=4, pady=4)
        self.price_var = tk.StringVar()
        self.price_entry = ttk.Entry(form, textvariable=self.price_var, width=12)
        self.price_entry.grid(row=0, column=3, padx=4, pady=4)

        ttk.Label(form, text="Date").grid(row=1, column=0, sticky="w", padx=4, pady=4)
        self.date_var = tk.StringVar(value=dt.date.today().isoformat())
        self.date_entry = ttk.Entry(form, textvariable=self.date_var, width=24)
        self.date_entry.grid(row=1, column=1, padx=4, pady=4)
        ttk.Label(form, text="(YYYY-MM-DD)", foreground="#888").grid(
            row=1, column=1, sticky="e", padx=4)

        ttk.Label(form, text="Category").grid(row=1, column=2, sticky="w", padx=4, pady=4)
        self.category_var = tk.StringVar()
        self.category_cb = ttk.Combobox(
            form, textvariable=self.category_var, state="readonly", width=18)
        self.category_cb.grid(row=1, column=3, padx=4, pady=4)

        btns = ttk.Frame(form)
        btns.grid(row=2, column=0, columnspan=4, sticky="w", pady=(8, 0))
        self.add_btn = ttk.Button(btns, text="Add (Ctrl+S)", command=self.add,
                                  underline=0)
        self.add_btn.pack(side="left", padx=(0, 6))
        self.update_btn = ttk.Button(
            btns, text="Update selected (Ctrl+S)", command=self.update,
            state="disabled")
        self.update_btn.pack(side="left", padx=6)
        self.delete_btn = ttk.Button(
            btns, text="Delete selected (Del)", command=self.delete,
            state="disabled")
        self.delete_btn.pack(side="left", padx=6)
        ttk.Button(btns, text="Clear (Esc)", command=self.clear_form).pack(
            side="left", padx=6)

    def _build_table(self):
        wrap = ttk.Frame(self)
        wrap.pack(fill="both", expand=True, pady=(10, 0))

        cols = ("id", "date", "store", "category", "price")
        self.tree = ttk.Treeview(wrap, columns=cols, show="headings", selectmode="browse")
        headings = {
            "id": "ID", "date": "Date", "store": "Store",
            "category": "Category", "price": "Price",
        }
        widths = {"id": 50, "date": 100, "store": 200, "category": 160, "price": 110}
        for c in cols:
            self.tree.heading(c, text=headings[c])
            anchor = "e" if c == "price" else "w"
            self.tree.column(c, width=widths[c], anchor=anchor)
        self.tree.pack(side="left", fill="both", expand=True)

        sb = ttk.Scrollbar(wrap, orient="vertical", command=self.tree.yview)
        sb.pack(side="right", fill="y")
        self.tree.configure(yscrollcommand=sb.set)
        self.tree.bind("<<TreeviewSelect>>", self.on_select)
        # Delete key removes the highlighted line; Enter loads it into the form.
        self.tree.bind("<Delete>", lambda _e: self.delete())
        self.tree.bind("<Return>", lambda _e: self.on_select())

    def _bind_keys(self):
        """Keyboard shortcuts scoped to the Data tab.

        Pressing Enter inside any form field adds or updates the line, so the
        whole flow can be driven without a mouse.
        """
        def _clear(_e):
            self.clear_form()
            return "break"  # don't fall through to the global Esc handler

        for entry in (self.store_entry, self.price_entry, self.date_entry):
            entry.bind("<Return>", lambda _e: self.save())
            entry.bind("<Escape>", _clear)
        # Enter on the category combobox behaves the same way.
        self.category_cb.bind("<Return>", lambda _e: self.save())

    # -- helpers ------------------------------------------------------------ #
    def refresh_categories(self):
        cats = self.db.categories()
        self._cat_by_name = {name: cid for cid, name in cats}
        names = [name for _, name in cats]
        self.category_cb["values"] = names
        if names and not self.category_var.get():
            self.category_var.set(names[0])

    def refresh(self):
        self.refresh_categories()
        for item in self.tree.get_children():
            self.tree.delete(item)
        for row in self.db.lines():
            self.tree.insert(
                "", "end", iid=str(row["id"]),
                values=(row["id"], row["date"], row["store"],
                        row["category_name"], gbp(row["price"])),
            )

    def _read_form(self):
        store = self.store_var.get().strip()
        price_s = self.price_var.get().strip().replace("£", "").replace(",", "")
        date_s = self.date_var.get().strip()
        cat_name = self.category_var.get()

        if not store:
            raise ValueError("Please enter a store.")
        try:
            price = round(float(price_s), 2)
        except ValueError:
            raise ValueError("Price must be a number, e.g. 12.99")
        if price < 0:
            raise ValueError("Price cannot be negative.")
        try:
            dt.date.fromisoformat(date_s)
        except ValueError:
            raise ValueError("Date must be in YYYY-MM-DD format.")
        if cat_name not in self._cat_by_name:
            raise ValueError("Please choose a category.")
        return store, price, date_s, self._cat_by_name[cat_name]

    # -- actions ------------------------------------------------------------ #
    def save(self):
        """Update the selected line, or add a new one (Ctrl+S / Enter)."""
        if self.selected_id is None:
            self.add()
        else:
            self.update()

    def focus_new(self):
        """Clear the form and focus the first field, ready for a new line."""
        self.clear_form()
        self.store_entry.focus_set()

    def add(self):
        try:
            store, price, date_s, cat_id = self._read_form()
        except ValueError as e:
            messagebox.showerror("Invalid input", str(e))
            return
        self.db.add_line(store, price, date_s, cat_id)
        self.clear_form()
        self.refresh()
        self.on_change()

    def update(self):
        if self.selected_id is None:
            return
        try:
            store, price, date_s, cat_id = self._read_form()
        except ValueError as e:
            messagebox.showerror("Invalid input", str(e))
            return
        self.db.update_line(self.selected_id, store, price, date_s, cat_id)
        self.clear_form()
        self.refresh()
        self.on_change()

    def delete(self):
        if self.selected_id is None:
            return
        if not messagebox.askyesno("Delete", "Delete the selected budget line?"):
            return
        self.db.delete_line(self.selected_id)
        self.clear_form()
        self.refresh()
        self.on_change()

    def on_select(self, _event=None):
        sel = self.tree.selection()
        if not sel:
            return
        self.selected_id = int(sel[0])
        vals = self.tree.item(sel[0], "values")
        # vals = (id, date, store, category, price)
        self.date_var.set(vals[1])
        self.store_var.set(vals[2])
        self.category_var.set(vals[3])
        # .replace, not .lstrip: gbp() puts the minus sign *before* the £ on
        # negative amounts ("-£12.34"), so lstrip("£") -- which only strips
        # from the very start -- would leave the £ in place and break the
        # float() parse in _read_form when re-editing such a line.
        self.price_var.set(vals[4].replace("£", "").replace(",", ""))
        self.update_btn.config(state="normal")
        self.delete_btn.config(state="normal")

    def clear_form(self):
        self.selected_id = None
        if self.tree.selection():
            self.tree.selection_remove(self.tree.selection())
        self.store_var.set("")
        self.price_var.set("")
        self.date_var.set(dt.date.today().isoformat())
        self.refresh_categories()
        self.update_btn.config(state="disabled")
        self.delete_btn.config(state="disabled")


# --------------------------------------------------------------------------- #
# Chart tab
# --------------------------------------------------------------------------- #
class ChartTab(ttk.Frame):
    """Column graph of totals grouped by category / year / month."""

    GROUPS = [("Category", "category"), ("Year", "year"), ("Month", "month")]

    def __init__(self, master, db, speaker=None):
        super().__init__(master, padding=10)
        self.db = db
        self.speaker = speaker or Speaker()

        top = ttk.Frame(self)
        top.pack(fill="x")
        ttk.Label(top, text="Group by:").pack(side="left")
        self.group_var = tk.StringVar(value="Category")
        cb = ttk.Combobox(
            top, textvariable=self.group_var, state="readonly",
            values=[label for label, _ in self.GROUPS], width=14)
        cb.pack(side="left", padx=8)
        cb.bind("<<ComboboxSelected>>", lambda _e: self.draw())
        ttk.Button(top, text="Refresh (Ctrl+R)", command=self.draw,
                   underline=0).pack(side="left")

        self.speak_btn = ttk.Button(
            top, text="Read chart aloud (Ctrl+T)", command=self.speak_chart)
        self.speak_btn.pack(side="left", padx=(12, 0))
        self.stop_btn = ttk.Button(
            top, text="Stop reading (Esc)", command=self.stop_speaking)
        self.stop_btn.pack(side="left", padx=6)
        if not self.speaker.available:
            self.speak_btn.state(["disabled"])
            self.stop_btn.state(["disabled"])

        self._build_speech_settings()

        self.canvas = tk.Canvas(self, bg="white", highlightthickness=2,
                                highlightbackground="#ccc",
                                highlightcolor="#4a90d9", takefocus=True)
        self.canvas.pack(fill="both", expand=True, pady=(10, 0))
        self.canvas.bind("<Configure>", lambda _e: self.draw())
        # The canvas can be focused with Tab; Enter or "r" reads it aloud so a
        # keyboard-only user never needs to reach for the button.
        self.canvas.bind("<Return>", lambda _e: self.speak_chart())
        self.canvas.bind("<r>", lambda _e: self.speak_chart())
        self.canvas.bind("<Escape>", lambda _e: self.stop_speaking())

        # A live text version of the chart: useful with a screen reader, and a
        # visible transcript of whatever is being read aloud.
        self.summary_var = tk.StringVar(
            value="Select a grouping and press Read chart aloud.")
        summary = ttk.Label(self, textvariable=self.summary_var,
                            wraplength=760, justify="left", foreground="#333")
        summary.pack(fill="x", pady=(8, 0))

    def _build_speech_settings(self):
        """Voice picker and speaking-rate slider for the read-aloud feature."""
        frame = ttk.Frame(self)
        frame.pack(fill="x", pady=(6, 0))

        voices = self.speaker.voices()
        self._voice_by_name = {name: vid for vid, name in voices}

        ttk.Label(frame, text="Voice:").pack(side="left")
        self.voice_var = tk.StringVar()
        self.voice_cb = ttk.Combobox(
            frame, textvariable=self.voice_var, state="readonly",
            width=30, values=[name for _, name in voices])
        self.voice_cb.pack(side="left", padx=(4, 16))
        self.voice_cb.bind("<<ComboboxSelected>>", self._on_voice)

        ttk.Label(frame, text="Speed:").pack(side="left")
        default_rate = self.speaker.default_rate
        self.rate_scale = ttk.Scale(
            frame, from_=80, to=320, orient="horizontal", length=180,
            command=self._on_rate)
        self.rate_scale.pack(side="left", padx=4)
        self.rate_label = ttk.Label(frame, width=9)
        self.rate_label.pack(side="left")

        # Create the label first: Scale.set() fires _on_rate, which needs it.
        self.rate_scale.set(default_rate)
        # Apply defaults up front so the very first utterance honours them.
        self.speaker.rate = int(default_rate)
        self._update_rate_label(default_rate)

        # Preselect the engine's current voice if we can identify it.
        name_by_id = {vid: name for vid, name in voices}
        current = name_by_id.get(self.speaker._default_voice)
        if current:
            self.voice_var.set(current)
            self.speaker.voice_id = self.speaker._default_voice
        elif voices:
            self.voice_var.set(voices[0][1])
            self.speaker.voice_id = voices[0][0]

        if not self.speaker.available:
            self.voice_cb.state(["disabled"])
            self.rate_scale.state(["disabled"])

    def _on_voice(self, _event=None):
        self.speaker.voice_id = self._voice_by_name.get(self.voice_var.get())

    def _on_rate(self, value):
        rate = int(float(value))
        self.speaker.rate = rate
        self._update_rate_label(rate)

    def _update_rate_label(self, rate):
        self.rate_label.config(text="{} wpm".format(int(float(rate))))

    def _group_key(self):
        label = self.group_var.get()
        return dict(self.GROUPS)[label]

    # -- accessibility: read the chart aloud -------------------------------- #
    def summary_text(self):
        """Build a spoken-language description of the current chart."""
        group_label = self.group_var.get().lower()
        data = self.db.totals_by(self._group_key())
        nonzero = [(label, val) for label, val in data if val and val > 0]

        if not nonzero:
            return ("Chart of total spend by {}. "
                    "There is no data to display yet.".format(group_label))

        total = sum(val for _, val in nonzero)
        ordered = sorted(nonzero, key=lambda t: t[1], reverse=True)
        top_label, top_val = ordered[0]

        parts = [
            "Total spend by {}.".format(group_label),
            "There {} {} {} with spending, totalling {}.".format(
                "is" if len(nonzero) == 1 else "are",
                len(nonzero),
                "entry" if len(nonzero) == 1 else "entries",
                spoken_gbp(total)),
            "The highest is {} at {}, {} percent of the total.".format(
                top_label, spoken_gbp(top_val),
                int(round(top_val / total * 100))),
        ]
        # Read every bar in the order it appears on screen.
        for label, val in nonzero:
            parts.append("{}: {}, {} percent.".format(
                label, spoken_gbp(val), int(round(val / total * 100))))
        return " ".join(parts)

    def speak_chart(self):
        """Announce the chart via text-to-speech and update the transcript."""
        text = self.summary_text()
        self.summary_var.set(text)
        if not self.speaker.available:
            self.summary_var.set(
                text + "\n\n(Install the 'pyttsx3' package to hear this "
                "read aloud: pip install pyttsx3)")
            return
        self.speaker.speak(text)

    def stop_speaking(self):
        self.speaker.stop()

    def draw(self):
        c = self.canvas
        c.delete("all")
        w = c.winfo_width()
        h = c.winfo_height()
        if w < 50 or h < 50:
            return

        data = self.db.totals_by(self._group_key())
        margin_left, margin_right = 90, 20
        margin_top, margin_bottom = 30, 60
        plot_w = w - margin_left - margin_right
        plot_h = h - margin_top - margin_bottom
        x0, y0 = margin_left, margin_top + plot_h  # bottom-left of plot area

        title = "Total spend by " + self.group_var.get().lower()
        c.create_text(w / 2, 15, text=title, font=("Helvetica", 13, "bold"))

        if not data or all(v <= 0 for _, v in data):
            c.create_text(w / 2, h / 2, text="No data to display",
                          fill="#888", font=("Helvetica", 12))
            return

        max_val = max(v for _, v in data)
        nice_max, step = self._nice_axis(max_val)

        # Y axis grid lines + GBP labels.
        n_ticks = int(round(nice_max / step)) if step else 1
        for i in range(n_ticks + 1):
            val = step * i
            y = y0 - (val / nice_max) * plot_h
            c.create_line(margin_left, y, w - margin_right, y, fill="#eee")
            c.create_text(margin_left - 8, y, text=gbp(val), anchor="e",
                          font=("Helvetica", 9))

        # Axes.
        c.create_line(x0, margin_top, x0, y0, fill="#333")
        c.create_line(x0, y0, w - margin_right, y0, fill="#333")

        # Bars.
        n = len(data)
        slot = plot_w / n
        bar_w = slot * 0.6
        for i, (label, val) in enumerate(data):
            cx = x0 + slot * i + slot / 2
            bar_h = (val / nice_max) * plot_h if nice_max else 0
            left = cx - bar_w / 2
            right = cx + bar_w / 2
            top = y0 - bar_h
            if val > 0:
                c.create_rectangle(left, top, right, y0,
                                   fill="#4a90d9", outline="#2f6da8")
                c.create_text(cx, top - 8, text=gbp(val),
                              font=("Helvetica", 8), fill="#333")
            # X label (rotated for readability when crowded).
            anchor = "n"
            if n > 8 or any(len(l) > 6 for l, _ in data):
                c.create_text(cx, y0 + 6, text=label, anchor="ne",
                              angle=35, font=("Helvetica", 8))
            else:
                c.create_text(cx, y0 + 6, text=label, anchor=anchor,
                              font=("Helvetica", 9))

    @staticmethod
    def _nice_axis(max_val):
        """Return (nice_max, step) giving a clean rounded Y axis."""
        if max_val <= 0:
            return 1.0, 1.0
        import math
        # Aim for ~5 ticks.
        raw_step = max_val / 5.0
        magnitude = 10 ** math.floor(math.log10(raw_step))
        for mult in (1, 2, 2.5, 5, 10):
            step = mult * magnitude
            if raw_step <= step:
                break
        nice_max = math.ceil(max_val / step) * step
        return nice_max, step


# --------------------------------------------------------------------------- #
# Report building & export
# --------------------------------------------------------------------------- #
def gather_report(db):
    """Collect everything the report needs from the database, once."""
    lines = db.lines()
    return {
        "generated": dt.datetime.now(),
        "source": db.config.description,
        "num_lines": len(lines),
        "total": sum(l["price"] for l in lines),
        "by_category": db.totals_by("category"),
        "by_year": db.totals_by("year"),
        "by_month": [(m, v) for m, v in db.totals_by("month") if v],
        "lines": lines,
    }


def _report_sections(data):
    """The three grouped sections shared by every export format."""
    return [
        ("Spending by category", data["by_category"]),
        ("Spending by year", data["by_year"]),
        ("Spending by month", data["by_month"]),
    ]


def report_speech_text(data):
    """A spoken-language version of the report, for the audio export."""
    if data["num_lines"] == 0:
        return "Budget report. No spending has been recorded yet."
    total = data["total"] or 1
    parts = [
        "Budget report.",
        "Total spend {} across {} budget line{}.".format(
            spoken_gbp(data["total"]), data["num_lines"],
            "" if data["num_lines"] == 1 else "s"),
    ]
    for title, rows in _report_sections(data):
        chunk = [title + "."]
        for label, val in rows:
            chunk.append("{}: {}, {} percent.".format(
                label, spoken_gbp(val), int(round(val / total * 100))))
        parts.append(" ".join(chunk))
    return " ".join(parts)


def write_pdf(data, path, include_lines=True):
    """Write the report as a PDF using fpdf2."""
    def t(s):  # core PDF fonts are latin-1; keep '£' but replace exotic chars
        return str(s).encode("latin-1", "replace").decode("latin-1")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 10, t("Budget Report"), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(110)
    pdf.cell(0, 6, t("Generated {}   |   Source: {}".format(
        data["generated"].strftime("%Y-%m-%d %H:%M"), data["source"])),
        new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0)
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 8, t("Summary"), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 6, t("Total spend: {}".format(gbp(data["total"]))),
             new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, t("Budget lines: {}".format(data["num_lines"])),
             new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    total = data["total"]
    for title, rows in _report_sections(data):
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 8, t(title), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 11)
        if not rows:
            pdf.cell(0, 6, t("(no data)"), new_x="LMARGIN", new_y="NEXT")
        for label, val in rows:
            share = " ({:.0f}%)".format(val / total * 100) if total else ""
            pdf.cell(120, 6, t(label))
            pdf.cell(0, 6, t(gbp(val) + share), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    if include_lines and data["lines"]:
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 8, t("All budget lines"), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "B", 10)
        for w, h in ((26, "Date"), (66, "Store"), (54, "Category")):
            pdf.cell(w, 6, t(h))
        pdf.cell(0, 6, t("Price"), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        for l in data["lines"]:
            pdf.cell(26, 6, t(l["date"]))
            pdf.cell(66, 6, t(l["store"][:34]))
            pdf.cell(54, 6, t(l["category_name"][:26]))
            pdf.cell(0, 6, t(gbp(l["price"])), new_x="LMARGIN", new_y="NEXT")

    pdf.output(path)


def write_docx(data, path, include_lines=True):
    """Write the report as a Word document using python-docx."""
    doc = docx.Document()
    doc.add_heading("Budget Report", 0)
    meta = doc.add_paragraph()
    run = meta.add_run("Generated {}   |   Source: {}".format(
        data["generated"].strftime("%Y-%m-%d %H:%M"), data["source"]))
    run.italic = True

    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Total spend: {}".format(gbp(data["total"])))
    doc.add_paragraph("Budget lines: {}".format(data["num_lines"]))

    total = data["total"]
    for title, rows in _report_sections(data):
        doc.add_heading(title, level=1)
        if not rows:
            doc.add_paragraph("(no data)")
            continue
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Item", "Amount", "Share"
        for label, val in rows:
            cells = table.add_row().cells
            cells[0].text = str(label)
            cells[1].text = gbp(val)
            cells[2].text = "{:.0f}%".format(val / total * 100) if total else "-"

    if include_lines and data["lines"]:
        doc.add_heading("All budget lines", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = "Date", "Store"
        hdr[2].text, hdr[3].text = "Category", "Price"
        for l in data["lines"]:
            cells = table.add_row().cells
            cells[0].text = l["date"]
            cells[1].text = l["store"]
            cells[2].text = l["category_name"]
            cells[3].text = gbp(l["price"])

    doc.save(path)


# --------------------------------------------------------------------------- #
# Report tab
# --------------------------------------------------------------------------- #
class ReportTab(ttk.Frame):
    """Save a spending report as a PDF, Word document, or audio file."""

    def __init__(self, master, db, speaker):
        super().__init__(master, padding=10)
        self.db = db
        self.speaker = speaker
        self._audio_result = queue.Queue()

        # (key, label, available?, hint-if-missing)
        self.formats = [
            ("pdf", "PDF document (.pdf)", FPDF is not None,
             "Install fpdf2:  pip install fpdf2"),
            ("docx", "Word document (.docx)", docx is not None,
             "Install python-docx:  pip install python-docx"),
            ("audio", "Audio file", self.speaker.available,
             "Install pyttsx3:  pip install pyttsx3"),
        ]

        ttk.Label(self, text="Save a spending report",
                  font=("Helvetica", 13, "bold")).pack(anchor="w")

        box = ttk.LabelFrame(self, text="Format", padding=10)
        box.pack(fill="x", pady=(10, 0))
        first_available = next((k for k, _, ok, _ in self.formats if ok), "pdf")
        self.format_var = tk.StringVar(value=first_available)
        for key, label, available, hint in self.formats:
            text = label if available else "{}  —  {}".format(label, hint)
            rb = ttk.Radiobutton(box, text=text, value=key,
                                 variable=self.format_var)
            rb.pack(anchor="w", pady=2)
            if not available:
                rb.state(["disabled"])

        self.include_lines_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(
            self, text="Include the full list of budget lines "
            "(PDF and Word only)",
            variable=self.include_lines_var).pack(anchor="w", pady=(10, 0))

        self.save_btn = ttk.Button(self, text="Save report… (Ctrl+S)",
                                   command=self.export)
        self.save_btn.pack(anchor="w", pady=(12, 0))
        if not any(ok for _, _, ok, _ in self.formats):
            self.save_btn.state(["disabled"])

        self.status_var = tk.StringVar(value="")
        ttk.Label(self, textvariable=self.status_var, wraplength=760,
                  justify="left", foreground="#333").pack(anchor="w",
                                                          pady=(12, 0))

    # DataTab/CategoriesTab expose save()/focus_new(); mirror save() so the
    # Ctrl+S shortcut works here too.
    def save(self):
        self.export()

    def _default_name(self, ext):
        return "budget-report-{}{}".format(
            dt.date.today().strftime("%Y%m%d"), ext)

    def export(self):
        fmt = self.format_var.get()
        if fmt == "audio":
            ext = ".aiff" if sys.platform == "darwin" else ".wav"
            filetypes = [("Audio file", "*" + ext), ("All files", "*.*")]
        elif fmt == "docx":
            ext, filetypes = ".docx", [("Word document", "*.docx")]
        else:
            ext, filetypes = ".pdf", [("PDF document", "*.pdf")]

        path = filedialog.asksaveasfilename(
            title="Save report", defaultextension=ext,
            initialfile=self._default_name(ext), filetypes=filetypes)
        if not path:
            return

        data = gather_report(self.db)
        try:
            if fmt == "pdf":
                write_pdf(data, path, self.include_lines_var.get())
            elif fmt == "docx":
                write_docx(data, path, self.include_lines_var.get())
            else:
                self._export_audio(data, path)
                return  # audio finishes asynchronously
        except Exception as e:
            messagebox.showerror("Export failed", str(e))
            self.status_var.set("Export failed: {}".format(e))
            return
        self._finish(path)

    def _export_audio(self, data, path):
        self.save_btn.state(["disabled"])
        self.status_var.set("Generating audio… this can take a few seconds.")
        started = self.speaker.save_to_file(
            report_speech_text(data), path,
            lambda err: self._audio_result.put((path, err)))
        if not started:
            self.save_btn.state(["!disabled"])
            self.status_var.set("Audio export is unavailable.")
            return
        self.after(150, self._poll_audio)

    def _poll_audio(self):
        try:
            path, err = self._audio_result.get_nowait()
        except queue.Empty:
            self.after(150, self._poll_audio)
            return
        self.save_btn.state(["!disabled"])
        if err:
            messagebox.showerror("Export failed", str(err))
            self.status_var.set("Audio export failed: {}".format(err))
        else:
            self._finish(path)

    def _finish(self, path):
        self.status_var.set("Saved report to:\n{}".format(path))
        messagebox.showinfo("Report saved", "Saved report to:\n{}".format(path))


# --------------------------------------------------------------------------- #
# Application
# --------------------------------------------------------------------------- #
class BudgetApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Budget Tracker")
        self.geometry("820x600")
        self.minsize(680, 480)

        self.db, db_warning = open_db()
        self.speaker = Speaker()

        # Show which database is in use, so a central-vs-local choice is visible.
        self.title("Budget Tracker — " + self.db.config.description)

        self.nb = ttk.Notebook(self)
        self.nb.pack(fill="both", expand=True)

        self.chart_tab = ChartTab(self.nb, self.db, speaker=self.speaker)
        self.data_tab = DataTab(self.nb, self.db, on_change=self._data_changed)
        self.categories_tab = CategoriesTab(
            self.nb, self.db, on_change=self._categories_changed)
        self.report_tab = ReportTab(self.nb, self.db, speaker=self.speaker)

        # `underline` gives each tab a mnemonic (Alt+D/C/G/R); enable_traversal
        # activates those plus Ctrl+Tab / Ctrl+Shift+Tab to cycle tabs.
        self.nb.add(self.data_tab, text="Data", underline=0)
        self.nb.add(self.chart_tab, text="Chart", underline=0)
        self.nb.add(self.categories_tab, text="Categories", underline=4)
        self.nb.add(self.report_tab, text="Report", underline=0)
        self.nb.enable_traversal()

        self._build_menu()
        self._bind_shortcuts()

        # If we asked for MariaDB but fell back to SQLite, tell the user why.
        if db_warning:
            self.after(200, lambda: messagebox.showwarning(
                "Database", db_warning))

        # Silent background check a couple seconds after startup; failures
        # and "already up to date" are not shown unless the user asked via
        # the Help menu (see _check_for_updates).
        self.after(2000, lambda: self._check_for_updates(interactive=False))

    # -- menu & shortcuts --------------------------------------------------- #
    def _build_menu(self):
        menubar = tk.Menu(self)

        filemenu = tk.Menu(menubar, tearoff=0)
        filemenu.add_command(label="Save report…", accelerator="Ctrl+4",
                             command=self._open_report)
        filemenu.add_separator()
        filemenu.add_command(label="Quit", accelerator="Ctrl+Q",
                             command=self.destroy)
        menubar.add_cascade(label="File", menu=filemenu, underline=0)

        editmenu = tk.Menu(menubar, tearoff=0)
        editmenu.add_command(label="New entry", accelerator="Ctrl+N",
                             command=self._editor_new)
        editmenu.add_command(label="Save entry", accelerator="Ctrl+S",
                             command=self._editor_save)
        editmenu.add_command(label="Delete entry", accelerator="Del",
                             command=self._editor_delete)
        menubar.add_cascade(label="Edit", menu=editmenu, underline=0)

        viewmenu = tk.Menu(menubar, tearoff=0)
        viewmenu.add_command(label="Data tab", accelerator="Ctrl+1",
                             command=lambda: self._select_tab(0))
        viewmenu.add_command(label="Chart tab", accelerator="Ctrl+2",
                             command=lambda: self._select_tab(1))
        viewmenu.add_command(label="Categories tab", accelerator="Ctrl+3",
                             command=lambda: self._select_tab(2))
        viewmenu.add_command(label="Report tab", accelerator="Ctrl+4",
                             command=lambda: self._select_tab(3))
        viewmenu.add_command(label="Refresh chart", accelerator="Ctrl+R",
                             command=self.chart_tab.draw)
        menubar.add_cascade(label="View", menu=viewmenu, underline=0)

        accmenu = tk.Menu(menubar, tearoff=0)
        state = "normal" if self.speaker.available else "disabled"
        accmenu.add_command(label="Read chart aloud", accelerator="Ctrl+T",
                            command=self._read_chart, state=state)
        accmenu.add_command(label="Stop reading", accelerator="Esc",
                            command=self.speaker.stop, state=state)
        menubar.add_cascade(label="Accessibility", menu=accmenu, underline=0)

        helpmenu = tk.Menu(menubar, tearoff=0)
        helpmenu.add_command(label="Keyboard shortcuts", accelerator="F1",
                             command=self._show_shortcuts)
        helpmenu.add_command(label="Check for Updates…",
                             command=lambda: self._check_for_updates(interactive=True))
        menubar.add_cascade(label="Help", menu=helpmenu, underline=0)

        self.config(menu=menubar)

    def _bind_shortcuts(self):
        # Bound on the toplevel with bind_all so they work whatever has focus.
        self.bind_all("<Control-q>", lambda _e: self.destroy())
        self.bind_all("<Control-n>", lambda _e: self._editor_new())
        self.bind_all("<Control-s>", lambda _e: self._editor_save())
        self.bind_all("<Control-r>", lambda _e: self.chart_tab.draw())
        self.bind_all("<Control-t>", lambda _e: self._read_chart())
        self.bind_all("<Control-Key-1>", lambda _e: self._select_tab(0))
        self.bind_all("<Control-Key-2>", lambda _e: self._select_tab(1))
        self.bind_all("<Control-Key-3>", lambda _e: self._select_tab(2))
        self.bind_all("<Control-Key-4>", lambda _e: self._select_tab(3))
        self.bind_all("<Escape>", lambda _e: self.speaker.stop())
        self.bind_all("<F1>", lambda _e: self._show_shortcuts())

    # -- cross-tab refresh -------------------------------------------------- #
    def _data_changed(self):
        """A budget line changed: update the chart and category totals."""
        self.chart_tab.draw()
        self.categories_tab.refresh()

    def _categories_changed(self):
        """A category changed: update the Data tab's picker and the chart."""
        self.data_tab.refresh()
        self.chart_tab.draw()

    # -- editing shortcuts act on whichever editable tab is showing --------- #
    def _current_editor(self):
        editors = {0: self.data_tab, 2: self.categories_tab,
                   3: self.report_tab}
        return editors.get(self.nb.index(self.nb.select()))

    def _editor_save(self):
        editor = self._current_editor()
        if editor is not None:  # Data.save / Categories.save / Report.export
            editor.save()
        return "break"

    def _editor_new(self):
        editor = self._current_editor()
        if editor is None or not hasattr(editor, "focus_new"):
            self._select_tab(0)  # e.g. Chart/Report -> start a new Data line
            editor = self.data_tab
        editor.focus_new()
        return "break"

    def _editor_delete(self):
        editor = self._current_editor()
        if editor is not None and hasattr(editor, "delete"):
            editor.delete()

    def _select_tab(self, index):
        self.nb.select(index)
        return "break"

    def _read_chart(self):
        """Switch to the chart, make sure it is current, then read it aloud."""
        self._select_tab(1)
        self.chart_tab.draw()
        self.chart_tab.speak_chart()
        return "break"

    def _open_report(self):
        self._select_tab(3)
        self.report_tab.save_btn.focus_set()
        return "break"

    def _show_shortcuts(self):
        messagebox.showinfo(
            "Keyboard shortcuts",
            "Navigation\n"
            "  Ctrl+1 / Alt+D   Data tab\n"
            "  Ctrl+2 / Alt+C   Chart tab\n"
            "  Ctrl+3 / Alt+G   Categories tab\n"
            "  Ctrl+4 / Alt+R   Report tab\n"
            "  Ctrl+Tab         Next tab\n"
            "  Tab / Shift+Tab  Move between controls\n\n"
            "Editing (Data & Categories tabs)\n"
            "  Ctrl+N           New entry (clear form)\n"
            "  Ctrl+S / Enter   Add or update the entry\n"
            "  Del              Delete the selected entry\n"
            "  Esc              Clear the form\n\n"
            "Report tab\n"
            "  Ctrl+S           Save the report\n\n"
            "Accessibility (Chart tab)\n"
            "  Ctrl+T           Read the chart aloud\n"
            "  Esc              Stop reading\n"
            "  Ctrl+R           Refresh the chart\n\n"
            "  F1               Show this help")

    # -- auto-update ---------------------------------------------------------- #
    # Worker threads only ever put results on a queue; every self.after() call
    # happens on the main thread (either the initiating call below, or a
    # poller that reschedules itself), matching the _audio_result/_poll_audio
    # pattern in ReportTab -- Tk widget methods are not documented as
    # thread-safe, so background threads must never call them directly.
    def _check_for_updates(self, interactive: bool):
        if getattr(self, "_update_busy", False):
            if interactive:
                messagebox.showinfo("Check for Updates",
                                     "An update check is already in progress.")
            return
        self._update_busy = True
        result_queue: queue.Queue = queue.Queue()
        threading.Thread(target=self._update_check_worker,
                          args=(interactive, result_queue), daemon=True).start()
        self.after(200, lambda: self._poll_update_check(result_queue))

    def _update_check_worker(self, interactive: bool, result_queue: queue.Queue):
        release = updater.check_latest_release(__version__)
        result_queue.put((release, interactive))

    def _poll_update_check(self, result_queue: queue.Queue):
        try:
            release, interactive = result_queue.get_nowait()
        except queue.Empty:
            self.after(200, lambda: self._poll_update_check(result_queue))
            return
        self._on_update_check_result(release, interactive)

    def _on_update_check_result(self, release, interactive: bool):
        if release is None:
            self._update_busy = False
            if interactive:
                messagebox.showinfo(
                    "Check for Updates",
                    f"You're up to date (version {__version__}).")
            return

        notes = release.notes or "No release notes provided."
        wants_update = messagebox.askyesno(
            "Update available",
            f"Version {release.version} is available (you have {__version__}).\n\n"
            f"{notes}\n\nUpdate now?")
        if not wants_update:
            self._update_busy = False
            return

        if not updater.is_frozen():
            self._update_busy = False
            messagebox.showinfo(
                "Update available",
                "Running from source, so this build can't update itself.\n\n"
                "Opening the release page in your browser…")
            webbrowser.open(updater.RELEASES_PAGE_URL)
            return

        if not release.asset_url:
            self._update_busy = False
            messagebox.showerror(
                "Update available",
                "No download is available for this platform yet.\n\n"
                "Opening the release page in your browser…")
            webbrowser.open(updater.RELEASES_PAGE_URL)
            return

        self._start_self_update(release)

    def _start_self_update(self, release):
        progress = tk.Toplevel(self)
        progress.title("Updating…")
        progress.resizable(False, False)
        tk.Label(progress, text=f"Downloading version {release.version}…",
                 padx=20, pady=12).pack()
        bar = ttk.Progressbar(progress, mode="indeterminate", length=240)
        bar.pack(padx=20, pady=(0, 16))
        bar.start(12)
        progress.transient(self)
        progress.grab_set()
        # There is no safe way to cancel mid-self-replace, so refuse the
        # window manager's close/Esc affordances rather than let the user
        # think they stopped something that's still running in the background.
        progress.protocol("WM_DELETE_WINDOW", lambda: None)

        error_queue: queue.Queue = queue.Queue()

        def worker():
            try:
                updater.perform_self_update(release)
                # perform_self_update calls os._exit(0) on success -- it
                # never returns here when it works.
            except Exception as exc:  # noqa: BLE001 - surface any failure
                error_queue.put(exc)

        threading.Thread(target=worker, daemon=True).start()
        self.after(200, lambda: self._poll_self_update(progress, error_queue))

    def _poll_self_update(self, progress, error_queue: queue.Queue):
        try:
            exc = error_queue.get_nowait()
        except queue.Empty:
            self.after(200, lambda: self._poll_self_update(progress, error_queue))
            return
        self._update_busy = False
        progress.destroy()
        messagebox.showerror(
            "Update failed",
            f"Couldn't install the update: {exc}\n\n"
            "Opening the release page in your browser so you can download it "
            "manually…")
        webbrowser.open(updater.RELEASES_PAGE_URL)


def main():
    reason = access_denied_reason(load_access_config())
    if reason:
        root = tk.Tk()
        root.withdraw()
        messagebox.showerror("Access denied", reason)
        root.destroy()
        sys.exit(1)
    app = BudgetApp()
    app.mainloop()


if __name__ == "__main__":
    main()
